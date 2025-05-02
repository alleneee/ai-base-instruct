"""文档处理器实现模块"""
import os
from typing import Dict, Any, List
import logging
import fitz  # PyMuPDF
import docx
import markitdown
from pathlib import Path
import datetime

from enterprise_kb.core.document_pipeline.base import DocumentProcessor, PipelineFactory
from enterprise_kb.core.config.settings import settings

logger = logging.getLogger(__name__)

@PipelineFactory.register_processor
class FileValidator(DocumentProcessor):
    """文件验证处理器"""
    
    SUPPORTED_TYPES = ['pdf', 'md', 'markdown', 'docx', 'txt', 'html']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """验证文件是否有效且可处理"""
        file_path = context.get('file_path')
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"文件不存在: {file_path}")
            
        file_type = context.get('file_type', '').lower()
        if not file_type or file_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件类型: {file_type}")
            
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > settings.MAX_FILE_SIZE_BYTES:
            raise ValueError(f"文件过大: {file_size} 字节，最大支持 {settings.MAX_FILE_SIZE_BYTES} 字节")
            
        # 更新上下文
        context['file_size'] = file_size
        context['validation_passed'] = True
        
        logger.info(f"文件验证通过: {file_path}, 类型: {file_type}, 大小: {file_size} 字节")
        return context


@PipelineFactory.register_processor
class MarkItDownProcessor(DocumentProcessor):
    """通用MarkItDown文档处理器，适用于多种文档格式"""
    
    # 支持更多格式，包括HTML、RTF等MarkItDown支持的格式
    SUPPORTED_TYPES = ['pdf', 'md', 'markdown', 'docx', 'txt', 'html', 'htm', 'rtf', 'odt', 'pptx']
    
    def __init__(self):
        """初始化MarkItDown处理器"""
        super().__init__()
        self.markitdown_converter = markitdown.MarkItDown(enable_plugins=True)
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """使用MarkItDown处理文档，转换为标准Markdown格式"""
        file_path = context.get('file_path')
        file_type = context.get('file_type', '').lower()
        
        if not file_path or not os.path.exists(file_path) or file_type not in self.SUPPORTED_TYPES:
            return context
            
        try:
            # 检查是否需要转换为Markdown
            if context.get('convert_to_markdown', True):
                logger.info(f"使用MarkItDown转换文档: {file_path}")
                
                # 使用MarkItDown转换文件
                try:
                    markdown_content = self.markitdown_converter.convert(file_path)
                    context['markdown_content'] = markdown_content
                    
                    # 同时提供纯文本内容，方便后续处理
                    if 'text_content' not in context:
                        context['text_content'] = markdown_content
                        
                    logger.info(f"文档转换为Markdown成功: {file_path}")
                    
                    # 提取文档结构信息
                    self._extract_document_structure(markdown_content, context)
                    
                except Exception as e:
                    logger.warning(f"MarkItDown转换失败，尝试使用备用方法: {str(e)}")
                    # 如果MarkItDown转换失败，使用原始处理流程
                    context['markdown_conversion_failed'] = True
            
        except Exception as e:
            logger.error(f"文档处理失败: {str(e)}")
            raise
            
        return context
    
    def _extract_document_structure(self, markdown_content: str, context: Dict[str, Any]):
        """从Markdown内容中提取文档结构信息"""
        import re
        
        # 提取标题
        headings = []
        heading_pattern = re.compile(r'^(#{1,6})\s+(.*?)$', re.MULTILINE)
        for match in heading_pattern.finditer(markdown_content):
            level = len(match.group(1))
            title = match.group(2).strip()
            headings.append((level, title))
        
        # 提取代码块
        code_blocks = []
        code_pattern = re.compile(r'```(\w*)\n(.*?)```', re.DOTALL)
        for match in code_pattern.finditer(markdown_content):
            language = match.group(1)
            code = match.group(2)
            code_blocks.append((language, code))
        
        # 提取图片
        images = []
        image_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
        for match in image_pattern.finditer(markdown_content):
            alt_text = match.group(1)
            image_path = match.group(2)
            images.append((alt_text, image_path))
        
        # 更新上下文
        context['document_structure'] = {
            'headings': headings,
            'code_blocks': code_blocks,
            'images': images
        }
        
        return context


@PipelineFactory.register_processor
class PDFProcessor(DocumentProcessor):
    """PDF文档处理器"""
    
    SUPPORTED_TYPES = ['pdf']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """处理PDF文档"""
        if context.get('file_type') != 'pdf':
            return context
            
        # 如果已经通过MarkItDown转换过，跳过额外处理
        if context.get('markdown_content') and not context.get('markdown_conversion_failed'):
            logger.info(f"PDF已通过MarkItDown处理，跳过额外处理")
            return context
            
        file_path = context.get('file_path')
        try:
            # 使用PyMuPDF提取文本
            doc = fitz.open(file_path)
            text_content = ""
            toc = doc.get_toc()  # 获取目录
            
            # 提取文本
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
                
            # 更新上下文
            context['text_content'] = text_content
            context['page_count'] = len(doc)
            context['toc'] = toc if toc else []
            
            # 如果需要转换为Markdown，但MarkItDown失败，使用备用方法
            if context.get('convert_to_markdown', True) and context.get('markdown_conversion_failed'):
                try:
                    # 尝试备用方法转换为Markdown
                    md_content = self._convert_to_basic_markdown(text_content, toc)
                    context['markdown_content'] = md_content
                    logger.info(f"PDF使用备用方法转换为Markdown成功: {file_path}")
                except Exception as e:
                    logger.error(f"PDF备用转换失败: {str(e)}")
                    
            logger.info(f"PDF处理完成: {file_path}, 页数: {len(doc)}")
            
        except Exception as e:
            logger.error(f"PDF处理失败: {str(e)}")
            raise
            
        return context
    
    def _convert_to_basic_markdown(self, text: str, toc: List) -> str:
        """简单地将文本转换为Markdown格式"""
        # 如果有目录，创建Markdown标题
        md_content = ""
        if toc:
            md_content += "# 文档目录\n\n"
            for level, title, page in toc:
                indent = "  " * (level - 1)
                md_content += f"{indent}- {title} (页码: {page})\n"
            md_content += "\n\n# 文档内容\n\n"
        
        # 添加文本内容，按段落分割
        paragraphs = text.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if para:
                md_content += para + "\n\n"
        
        return md_content


@PipelineFactory.register_processor
class DocxProcessor(DocumentProcessor):
    """Word文档处理器"""
    
    SUPPORTED_TYPES = ['docx']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """处理Word文档"""
        if context.get('file_type') != 'docx':
            return context
            
        # 如果已经通过MarkItDown转换过，跳过额外处理
        if context.get('markdown_content') and not context.get('markdown_conversion_failed'):
            logger.info(f"Word文档已通过MarkItDown处理，跳过额外处理")
            return context
            
        file_path = context.get('file_path')
        try:
            # 使用python-docx提取文本
            doc = docx.Document(file_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])
            
            # 更新上下文
            context['text_content'] = text_content
            context['paragraph_count'] = len(doc.paragraphs)
            
            # 如果需要转换为Markdown，可以在这里调用MarkItDown
            if context.get('convert_to_markdown', True):
                context['markdown_content'] = self._convert_to_markdown(file_path)
                
            logger.info(f"Word文档处理完成: {file_path}, 段落数: {len(doc.paragraphs)}")
            
        except Exception as e:
            logger.error(f"Word文档处理失败: {str(e)}")
            raise
            
        return context
        
    def _convert_to_markdown(self, file_path: str) -> str:
        """
        将Word文档转换为Markdown
        
        Args:
            file_path: Word文件路径
            
        Returns:
            Markdown内容
        """
        try:
            # 使用MarkItDown转换
            md_content = markitdown.markitdown(file_path)
            return md_content
        except Exception as e:
            logger.error(f"Word转Markdown失败: {str(e)}")
            raise


@PipelineFactory.register_processor
class MarkdownProcessor(DocumentProcessor):
    """Markdown文档处理器"""
    
    SUPPORTED_TYPES = ['md', 'markdown']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """处理Markdown文档"""
        if context.get('file_type') not in ['md', 'markdown']:
            return context
            
        file_path = context.get('file_path')
        try:
            # 读取Markdown文件
            with open(file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
                
            # 更新上下文
            context['text_content'] = markdown_content
            context['markdown_content'] = markdown_content  # 已经是Markdown格式
            
            logger.info(f"Markdown文档处理完成: {file_path}")
            
        except Exception as e:
            logger.error(f"Markdown处理失败: {str(e)}")
            raise
            
        return context


@PipelineFactory.register_processor
class TextProcessor(DocumentProcessor):
    """文本文件处理器"""
    
    SUPPORTED_TYPES = ['txt']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """处理纯文本文件"""
        if context.get('file_type') != 'txt':
            return context
            
        file_path = context.get('file_path')
        try:
            # 读取文本文件
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
                
            # 更新上下文
            context['text_content'] = text_content
            
            logger.info(f"文本文件处理完成: {file_path}")
            
        except Exception as e:
            logger.error(f"文本处理失败: {str(e)}")
            raise
            
        return context


@PipelineFactory.register_processor
class ChunkingProcessor(DocumentProcessor):
    """文档分块处理器"""
    
    SUPPORTED_TYPES = ['pdf', 'md', 'markdown', 'docx', 'txt', 'html']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """根据文档类型选择合适的分块策略"""
        text_content = context.get('text_content')
        if not text_content:
            return context
            
        file_type = context.get('file_type')
        
        # 根据文件类型选择分块策略
        if file_type in ['md', 'markdown']:
            chunks = self._chunk_markdown(text_content)
        elif file_type == 'pdf':
            chunks = self._chunk_pdf(text_content)
        else:
            chunks = self._chunk_text(text_content)
            
        # 更新上下文
        context['chunks'] = chunks
        context['chunk_count'] = len(chunks)
        
        logger.info(f"文档分块完成，共 {len(chunks)} 个块")
        return context
        
    def _chunk_markdown(self, text: str) -> List[str]:
        """
        按Markdown结构分块
        
        Args:
            text: Markdown文本
            
        Returns:
            文本块列表
        """
        # 解析Markdown文档结构
        structure = self._parse_markdown_structure(text)
        
        # 基于结构进行智能分块
        return self._structure_based_chunking(structure)
    
    def _parse_markdown_structure(self, text: str) -> Dict[str, Any]:
        """
        解析Markdown文档结构
        
        Args:
            text: Markdown文本
            
        Returns:
            文档结构字典，包含标题层级树和特殊内容块
        """
        import re
        
        # 结构对象
        structure = {
            'headings': [],  # 标题及其层级
            'sections': [],  # 文档区块
            'special_blocks': [],  # 特殊内容块（代码块、表格等）
            'references': []  # 新增: 引用关系
        }
        
        # 分割成行
        lines = text.split('\n')
        
        # 标题正则表达式
        heading_pattern = re.compile(r'^(#{1,6})\s+(.*?)$')
        
        # 代码块识别
        in_code_block = False
        code_block_start = 0
        code_block_language = ""
        
        # 表格识别
        table_rows = []
        in_table = False
        table_start = 0
        
        # 列表识别
        list_pattern = re.compile(r'^\s*(?:[*+-]|\d+\.)\s')
        in_list = False
        list_items = []
        list_start = 0
        
        # 引用关系识别模式 - 新增
        # 1. Markdown链接格式 [text](url)
        link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
        # 2. 引用格式 [ref: xxx] 或 [[xxx]]
        ref_pattern = re.compile(r'\[ref:([^\]]+)\]|\[\[([^\]]+)\]\]')
        # 3. 中文引用格式 如「参见：xxx」
        cn_ref_pattern = re.compile(r'[「『]参见[:：]([^」』]+)[」』]')
        
        current_section = {
            'level': 0,  # 顶级
            'title': "",
            'content': "",
            'start_line': 0,
            'end_line': 0,
            'subsections': [],
            'refs': []  # 新增: 当前段落中的引用
        }
        
        section_stack = [current_section]
        
        # 标题ID映射，用于引用解析
        heading_id_map = {}  # 标题文本到ID的映射
        
        # 遍历每一行
        for i, line in enumerate(lines):
            # 检查是否是标题行
            heading_match = heading_pattern.match(line)
            
            # 处理代码块
            if line.strip().startswith("```"):
                if not in_code_block:
                    # 代码块开始
                    in_code_block = True
                    code_block_start = i
                    # 提取语言
                    code_block_language = line.strip()[3:].strip()
                else:
                    # 代码块结束
                    in_code_block = False
                    structure['special_blocks'].append({
                        'type': 'code',
                        'language': code_block_language,
                        'start_line': code_block_start,
                        'end_line': i,
                        'content': '\n'.join(lines[code_block_start+1:i])
                    })
                continue
                
            # 在代码块内，跳过结构分析
            if in_code_block:
                continue
                
            # 处理表格
            if line.strip().startswith("|") and line.strip().endswith("|"):
                if not in_table:
                    # 表格开始
                    in_table = True
                    table_start = i
                    table_rows = [line]
                else:
                    # 继续表格
                    table_rows.append(line)
            elif in_table:
                # 表格结束
                in_table = False
                if len(table_rows) > 1:  # 至少包含表头和分隔行
                    structure['special_blocks'].append({
                        'type': 'table',
                        'start_line': table_start,
                        'end_line': i-1,
                        'content': '\n'.join(table_rows)
                    })
                table_rows = []
                
            # 处理列表
            list_match = list_pattern.match(line)
            if list_match:
                if not in_list:
                    # 列表开始
                    in_list = True
                    list_start = i
                    list_items = [line]
                else:
                    # 继续列表
                    list_items.append(line)
            elif in_list and line.strip() == "":
                # 暂时保持列表状态，可能是列表项之间的空行
                list_items.append(line)
            elif in_list:
                # 列表结束
                in_list = False
                structure['special_blocks'].append({
                    'type': 'list',
                    'start_line': list_start,
                    'end_line': i-1,
                    'content': '\n'.join(list_items)
                })
                list_items = []
                
            # 处理引用关系 - 新增
            # 检查Markdown链接
            for match in link_pattern.finditer(line):
                link_text = match.group(1)
                link_url = match.group(2)
                
                # 判断是否是内部引用(以#开头的锚点)
                if link_url.startswith('#'):
                    ref_id = link_url[1:]  # 去掉#号
                    structure['references'].append({
                        'type': 'internal',
                        'source': {
                            'line': i,
                            'section': current_section['title'],
                            'context': line
                        },
                        'target': ref_id,
                        'text': link_text
                    })
                    
                    # 添加到当前段落的引用中
                    current_section['refs'].append({
                        'type': 'internal',
                        'target': ref_id,
                        'text': link_text
                    })
            
            # 检查引用格式 [ref: xxx] 或 [[xxx]]
            for match in ref_pattern.finditer(line):
                ref_text = match.group(1) if match.group(1) else match.group(2)
                structure['references'].append({
                    'type': 'reference',
                    'source': {
                        'line': i,
                        'section': current_section['title'],
                        'context': line
                    },
                    'target': ref_text,
                    'text': ref_text
                })
                
                # 添加到当前段落的引用中
                current_section['refs'].append({
                    'type': 'reference',
                    'target': ref_text,
                    'text': ref_text
                })
            
            # 检查中文引用格式
            for match in cn_ref_pattern.finditer(line):
                ref_text = match.group(1)
                structure['references'].append({
                    'type': 'cn_reference',
                    'source': {
                        'line': i,
                        'section': current_section['title'],
                        'context': line
                    },
                    'target': ref_text,
                    'text': ref_text
                })
                
                # 添加到当前段落的引用中
                current_section['refs'].append({
                    'type': 'cn_reference',
                    'target': ref_text,
                    'text': ref_text
                })
                
            # 处理标题和段落结构
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                # 生成标题ID (用于引用解析)
                heading_id = self._generate_heading_id(title)
                heading_id_map[title.lower()] = heading_id
                
                # 保存当前段落
                if current_section['content'].strip():
                    current_section['end_line'] = i - 1
                
                # 创建新的段落
                new_section = {
                    'level': level,
                    'title': title,
                    'content': "",
                    'start_line': i,
                    'end_line': 0,
                    'subsections': [],
                    'refs': [],  # 引用列表
                    'heading_id': heading_id  # 标题ID
                }
                
                # 更新标题列表
                structure['headings'].append((level, title, i, heading_id))
                
                # 调整段落层级
                while len(section_stack) > 1 and section_stack[-1]['level'] >= level:
                    section_stack.pop()
                
                # 添加到当前段落
                section_stack[-1]['subsections'].append(new_section)
                section_stack.append(new_section)
                current_section = new_section
            else:
                # 普通内容行，添加到当前段落
                if current_section['content']:
                    current_section['content'] += '\n' + line
                else:
                    current_section['content'] = line
        
        # 处理最后一个段落
        if current_section['content'].strip():
            current_section['end_line'] = len(lines) - 1
        
        # 处理可能未关闭的特殊块
        if in_code_block:
            structure['special_blocks'].append({
                'type': 'code',
                'language': code_block_language,
                'start_line': code_block_start,
                'end_line': len(lines) - 1,
                'content': '\n'.join(lines[code_block_start+1:])
            })
            
        if in_table and len(table_rows) > 1:
            structure['special_blocks'].append({
                'type': 'table',
                'start_line': table_start,
                'end_line': len(lines) - 1,
                'content': '\n'.join(table_rows)
            })
            
        if in_list:
            structure['special_blocks'].append({
                'type': 'list',
                'start_line': list_start,
                'end_line': len(lines) - 1,
                'content': '\n'.join(list_items)
            })
        
        # 将层级结构添加到结果
        structure['sections'] = section_stack[0]['subsections']
        
        # 引用解析 - 将引用与实际标题关联起来
        self._resolve_references(structure, heading_id_map)
        
        return structure
    
    def _generate_heading_id(self, heading_text: str) -> str:
        """
        为标题生成一个ID，用于引用
        
        Args:
            heading_text: 标题文本
            
        Returns:
            标题ID
        """
        # 简化版的ID生成，将标题转为小写并替换非字母数字为连字符
        import re
        # 移除HTML标签
        text = re.sub(r'<[^>]+>', '', heading_text)
        # 转为小写
        text = text.lower()
        # 替换非字母数字字符为连字符
        text = re.sub(r'[^a-z0-9\u4e00-\u9fa5]+', '-', text)
        # 移除开头和结尾的连字符
        text = text.strip('-')
        return text
    
    def _resolve_references(self, structure: Dict[str, Any], heading_id_map: Dict[str, str]):
        """
        解析引用关系，将引用与实际标题关联
        
        Args:
            structure: 文档结构
            heading_id_map: 标题到ID的映射
        """
        # 反向映射：ID到标题索引
        id_to_heading = {}
        for i, (level, title, line, heading_id) in enumerate(structure['headings']):
            id_to_heading[heading_id] = i
            id_to_heading[title.lower()] = i  # 同时支持按标题文本查找
        
        # 解析每个引用
        resolved_refs = []
        for ref in structure['references']:
            target = ref['target']
            
            # 尝试查找目标标题
            target_id = None
            target_index = -1
            
            # 1. 直接匹配ID
            if target in id_to_heading:
                target_index = id_to_heading[target]
            # 2. 匹配标题文本
            elif target.lower() in id_to_heading:
                target_index = id_to_heading[target.lower()]
            # 3. 通过标题文本获取ID
            elif target.lower() in heading_id_map:
                target_id = heading_id_map[target.lower()]
                if target_id in id_to_heading:
                    target_index = id_to_heading[target_id]
            
            # 找到匹配的标题
            if target_index >= 0:
                level, title, line, heading_id = structure['headings'][target_index]
                resolved_refs.append({
                    'source': ref['source'],
                    'target': {
                        'type': 'heading',
                        'title': title,
                        'level': level,
                        'line': line,
                        'id': heading_id
                    },
                    'text': ref['text'],
                    'resolved': True
                })
            else:
                # 未找到匹配的标题
                resolved_refs.append({
                    'source': ref['source'],
                    'target': {
                        'type': 'unknown',
                        'text': target
                    },
                    'text': ref['text'],
                    'resolved': False
                })
        
        # 用解析后的引用替换原引用列表
        structure['references'] = resolved_refs
    
    def _structure_based_chunking(self, structure: Dict[str, Any]) -> List[str]:
        """
        基于文档结构进行智能分块
        
        Args:
            structure: 文档结构字典
            
        Returns:
            分块列表
        """
        chunks = []
        max_chunk_size = settings.MAX_CHUNK_SIZE
        
        # 处理特殊块
        special_blocks_map = {}
        for block in structure['special_blocks']:
            # 创建行号到特殊块的映射
            for line_num in range(block['start_line'], block['end_line'] + 1):
                special_blocks_map[line_num] = block
        
        # 引用映射: 标题ID到被引用计数
        ref_count_map = {}
        ref_source_map = {}  # 源引用到目标引用的映射
        
        # 统计引用次数
        for ref in structure['references']:
            if ref.get('resolved', False):
                target_id = ref['target']['id']
                ref_count_map[target_id] = ref_count_map.get(target_id, 0) + 1
                
                # 记录源位置
                source_section = ref['source']['section']
                if source_section not in ref_source_map:
                    ref_source_map[source_section] = []
                ref_source_map[source_section].append(ref)
        
        # 递归处理段落
        def process_section(section, parent_titles=None, is_referenced=False):
            nonlocal chunks
            
            if parent_titles is None:
                parent_titles = []
            
            # 创建当前块的标题上下文
            current_titles = parent_titles.copy()
            if section.get('title'):
                current_titles.append(section['title'])
            
            # 构建完整标题路径
            title_path = " > ".join(current_titles)
            title_prefix = f"{title_path}\n\n" if title_path else ""
            
            # 检查当前段落是否被引用
            section_is_referenced = is_referenced
            if not section_is_referenced and 'heading_id' in section:
                section_is_referenced = ref_count_map.get(section['heading_id'], 0) > 0
            
            # 处理当前段落内容
            content = section.get('content', '').strip()
            if content:
                # 检查当前段落是否包含引用
                has_outgoing_refs = len(section.get('refs', [])) > 0
                
                # 计算元数据部分
                metadata = ""
                if section_is_referenced or has_outgoing_refs:
                    metadata_parts = []
                    if section_is_referenced:
                        ref_count = ref_count_map.get(section.get('heading_id', ''), 0)
                        metadata_parts.append(f"[被引用次数: {ref_count}]")
                    
                    if has_outgoing_refs:
                        outgoing_count = len(section.get('refs', []))
                        metadata_parts.append(f"[外部引用数: {outgoing_count}]")
                    
                    if metadata_parts:
                        metadata = " ".join(metadata_parts) + "\n"
                
                # 完整内容 = 标题路径 + 元数据 + 内容
                full_content = title_prefix + metadata + content
                
                # 检查内容大小
                if len(full_content) <= max_chunk_size:
                    # 直接作为一个块
                    chunks.append(full_content)
                else:
                    # 需要进一步分割，使用更智能的结构感知分割
                    content_chunks = self._structure_aware_split(content)
                    for i, chunk in enumerate(content_chunks):
                        # 只给第一个块添加完整元数据
                        if i == 0:
                            chunks.append(title_prefix + metadata + chunk)
                        else:
                            # 其他块只添加标题路径，避免重复
                            chunks.append(title_prefix + chunk)
            
            # 递归处理子段落
            for subsection in section.get('subsections', []):
                # 检查该子段落是否被引用
                subsection_is_referenced = section_is_referenced
                if not subsection_is_referenced and 'heading_id' in subsection:
                    subsection_is_referenced = ref_count_map.get(subsection['heading_id'], 0) > 0
                
                process_section(subsection, current_titles, subsection_is_referenced)
        
        # 处理文档的顶级段落
        for section in structure.get('sections', []):
            process_section(section)
        
        # 后处理：处理引用关系
        processed_chunks = self._process_reference_contexts(chunks, structure['references'])
        
        # 如果没有结构化的块，则回退到基本分块
        if not processed_chunks:
            return self._chunk_by_headings(structure.get('raw_text', ''), ['# ', '## ', '### '])
        
        return processed_chunks
        
    def _process_reference_contexts(self, chunks: List[str], references: List[Dict]) -> List[str]:
        """
        处理引用上下文，确保相互引用的内容在语义上保持连贯
        
        Args:
            chunks: 初步分块结果
            references: 引用关系列表
            
        Returns:
            处理后的分块
        """
        # 如果没有引用，直接返回原始分块
        if not references:
            return chunks
            
        # 引用计数器 - 用于识别重要引用
        ref_importance = {}
        
        # 统计每个引用的重要性
        for ref in references:
            if ref.get('resolved', False):
                target_id = ref['target']['id']
                ref_importance[target_id] = ref_importance.get(target_id, 0) + 1
        
        # 创建重要引用的上下文映射
        ref_contexts = {}
        important_threshold = 2  # 被引用次数超过2次视为重要引用
        
        for target_id, count in ref_importance.items():
            if count >= important_threshold:
                # 为重要引用构建上下文摘要
                context = self._extract_reference_context(target_id, references)
                if context:
                    ref_contexts[target_id] = context
        
        # 处理每个分块，为包含重要引用的分块添加上下文
        result_chunks = []
        
        for chunk in chunks:
            modified_chunk = chunk
            
            # 检查分块中是否包含重要引用的标记
            for target_id, context in ref_contexts.items():
                # 检查分块中是否包含对该引用的引用
                for ref_type in ['[ref:', '[[', '「参见']:
                    if ref_type in chunk and target_id in chunk:
                        # 在分块末尾添加引用上下文
                        if not modified_chunk.endswith(context):
                            modified_chunk += f"\n\n引用上下文: {context}"
                        break
            
            result_chunks.append(modified_chunk)
            
        return result_chunks
        
    def _extract_reference_context(self, target_id: str, references: List[Dict]) -> str:
        """
        提取引用目标的上下文摘要
        
        Args:
            target_id: 引用目标ID
            references: 引用关系列表
            
        Returns:
            上下文摘要
        """
        # 找到匹配的引用
        target_refs = [ref for ref in references if ref.get('resolved', False) and ref['target']['id'] == target_id]
        
        if not target_refs:
            return ""
            
        # 使用第一个匹配的引用
        ref = target_refs[0]
        target = ref['target']
        
        # 提取上下文
        context = f"{target['title']} (级别: {target['level']})"
        
        return context
    
    def _structure_aware_split(self, text: str) -> List[str]:
        """
        结构感知的文本分割
        识别段落、列表、代码块等结构，更智能地进行分割
        
        Args:
            text: 要分割的文本
            
        Returns:
            分割后的文本块
        """
        import re
        
        # 识别文本中的结构元素
        structures = []
        
        # 识别代码块
        code_pattern = re.compile(r'```[\w]*\n.*?```', re.DOTALL)
        for match in code_pattern.finditer(text):
            structures.append({
                'type': 'code',
                'start': match.start(),
                'end': match.end(),
                'content': match.group()
            })
        
        # 识别列表
        list_pattern = re.compile(r'(?:(?:\n|^)\s*(?:[*+-]|\d+\.)\s+.*)+', re.DOTALL)
        for match in list_pattern.finditer(text):
            # 检查是否与已识别的结构重叠
            overlap = False
            for s in structures:
                if (match.start() < s['end'] and match.end() > s['start']):
                    overlap = True
                    break
            if not overlap:
                structures.append({
                    'type': 'list',
                    'start': match.start(),
                    'end': match.end(),
                    'content': match.group()
                })
        
        # 识别表格
        table_pattern = re.compile(r'(?:\|.*\|(?:\n|$))+', re.DOTALL)
        for match in table_pattern.finditer(text):
            # 检查是否与已识别的结构重叠
            overlap = False
            for s in structures:
                if (match.start() < s['end'] and match.end() > s['start']):
                    overlap = True
                    break
            if not overlap:
                structures.append({
                    'type': 'table',
                    'start': match.start(),
                    'end': match.end(),
                    'content': match.group()
                })
        
        # 按开始位置排序结构
        structures.sort(key=lambda x: x['start'])
        
        # 如果没有识别到结构，使用递归重叠分块
        if not structures:
            return self._recursive_overlap_chunking(text, settings.MAX_CHUNK_SIZE, 100)
        
        # 构建分块
        chunks = []
        last_end = 0
        
        for structure in structures:
            # 处理结构前的文本
            if structure['start'] > last_end:
                pre_text = text[last_end:structure['start']].strip()
                if pre_text:
                    # 对普通文本使用递归重叠分块
                    pre_chunks = self._recursive_overlap_chunking(pre_text, settings.MAX_CHUNK_SIZE, 100)
                    chunks.extend(pre_chunks)
            
            # 处理结构本身
            structure_content = structure['content']
            
            # 如果结构内容超过块大小，尝试智能分割
            if len(structure_content) > settings.MAX_CHUNK_SIZE:
                if structure['type'] == 'code':
                    # 代码块尝试按函数/类分割
                    code_chunks = self._split_code_block(structure_content)
                    chunks.extend(code_chunks)
                elif structure['type'] == 'list':
                    # 列表按项分割
                    list_chunks = self._split_list(structure_content)
                    chunks.extend(list_chunks)
                elif structure['type'] == 'table':
                    # 表格可能需要特殊处理，这里简单处理
                    table_chunks = self._split_by_size(structure_content, settings.MAX_CHUNK_SIZE, 50)
                    chunks.extend(table_chunks)
            else:
                # 结构内容不大，直接作为一个块
                chunks.append(structure_content)
            
            last_end = structure['end']
        
        # 处理最后一部分文本
        if last_end < len(text):
            post_text = text[last_end:].strip()
            if post_text:
                post_chunks = self._recursive_overlap_chunking(post_text, settings.MAX_CHUNK_SIZE, 100)
                chunks.extend(post_chunks)
        
        return chunks
    
    def _recursive_overlap_chunking(self, text: str, max_size: int, overlap_size: int) -> List[str]:
        """
        递归重叠分块 - 借鉴LlamaIndex的设计思想
        
        当文本较长时，先分成大块，再递归地处理每个大块，保持上下文连贯性
        
        Args:
            text: 文本内容
            max_size: 最大块大小
            overlap_size: 重叠大小
            
        Returns:
            分块列表
        """
        # 空文本或小文本直接返回
        if not text or len(text) <= max_size:
            return [text] if text else []
        
        chunks = []
        
        # 按段落分割文本
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            
            # 跳过空段落
            if not para:
                continue
                
            # 如果段落本身超过最大大小
            if len(para) > max_size:
                # 先保存当前累积的块
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                
                # 大段落递归处理
                if len(para) > max_size * 2:  # 非常大的段落
                    # 先按句子分割
                    sentences = self._split_into_sentences(para)
                    # 再对句子应用递归重叠
                    para_chunks = self._chunk_sentences_with_overlap(sentences, max_size, overlap_size)
                else:
                    # 中等大小的段落，直接应用句子分割
                    para_chunks = self._split_by_sentences(para)
                
                chunks.extend(para_chunks)
                continue
            
            # 如果当前块 + 新段落超过大小限制
            if current_chunk and len(current_chunk) + len(para) + 2 > max_size:
                # 保存当前块
                chunks.append(current_chunk)
                
                # 新段落作为新块的开始，添加重叠
                if overlap_size > 0 and len(current_chunk) > overlap_size:
                    # 从当前块的末尾提取重叠部分
                    overlap_text = self._extract_semantic_overlap(current_chunk, overlap_size)
                    current_chunk = overlap_text + "\n\n" + para
                else:
                    current_chunk = para
            else:
                # 添加到当前块
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
        
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        将文本分割成句子
        
        Args:
            text: 文本内容
            
        Returns:
            句子列表
        """
        import re
        
        # 句子结束标记正则表达式 - 支持中英文
        sentence_end = re.compile(r'([.!?。！？；;]\s*)')
        
        # 分割成句子
        sentences = sentence_end.split(text)
        
        # 合并句子和标点
        result = []
        for i in range(0, len(sentences)-1, 2):
            if i+1 < len(sentences):
                result.append(sentences[i] + sentences[i+1])
            else:
                result.append(sentences[i])
        
        # 处理最后一个元素
        if len(sentences) % 2 == 1:
            result.append(sentences[-1])
        
        return [s for s in result if s.strip()]
    
    def _chunk_sentences_with_overlap(self, sentences: List[str], max_size: int, overlap_size: int) -> List[str]:
        """
        将句子列表分块，并保持重叠
        
        Args:
            sentences: 句子列表
            max_size: 最大块大小
            overlap_size: 重叠大小
            
        Returns:
            分块列表
        """
        chunks = []
        current_chunk = ""
        last_added_index = -1
        
        for i, sentence in enumerate(sentences):
            # 如果单个句子就超过了最大大小
            if len(sentence) > max_size:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                
                # 分割大句子
                sentence_parts = self._split_by_size(sentence, max_size, 0)
                chunks.extend(sentence_parts)
                last_added_index = i
                continue
            
            # 如果当前块 + 新句子超过大小限制
            if current_chunk and len(current_chunk) + len(sentence) + 1 > max_size:
                chunks.append(current_chunk)
                
                # 确定重叠句子的开始索引
                overlap_start = max(last_added_index - 2, 0)  # 尝试包含前两个句子
                current_chunk = " ".join(sentences[overlap_start:i]) + " " + sentence
                last_added_index = i
            else:
                # 添加到当前块
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                last_added_index = i
        
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _extract_semantic_overlap(self, text: str, overlap_size: int) -> str:
        """
        从文本中提取语义上有意义的重叠部分
        
        尝试提取完整句子或段落，而不是简单截断
        
        Args:
            text: 源文本
            overlap_size: 期望的重叠大小
            
        Returns:
            重叠文本
        """
        if len(text) <= overlap_size:
            return text
        
        # 尝试在段落边界分割
        paragraphs = text.split('\n\n')
        
        # 从后向前构建重叠文本
        overlap_text = ""
        for para in reversed(paragraphs):
            if len(para) + len(overlap_text) + 2 <= overlap_size:
                if overlap_text:
                    overlap_text = para + "\n\n" + overlap_text
                else:
                    overlap_text = para
            else:
                # 当前段落会超过重叠大小，需要进一步拆分
                remaining_size = overlap_size - len(overlap_text)
                if remaining_size > 0:
                    # 尝试在句子边界分割
                    sentences = self._split_into_sentences(para)
                    sentence_overlap = ""
                    
                    for sentence in reversed(sentences):
                        if len(sentence) + len(sentence_overlap) + 1 <= remaining_size:
                            if sentence_overlap:
                                sentence_overlap = sentence + " " + sentence_overlap
                            else:
                                sentence_overlap = sentence
                        else:
                            break
                    
                    if sentence_overlap:
                        if overlap_text:
                            overlap_text = sentence_overlap + "\n\n" + overlap_text
                        else:
                            overlap_text = sentence_overlap
                
                break
        
        # 如果无法在段落或句子边界获得足够大小的重叠，使用直接截断
        if not overlap_text or len(overlap_text) < overlap_size / 2:
            overlap_text = text[-overlap_size:]
        
        return overlap_text
    
    def _split_by_size(self, text: str, max_size: int, overlap_size: int) -> List[str]:
        """
        按固定大小分割文本，保持重叠
        
        Args:
            text: 文本内容
            max_size: 最大块大小
            overlap_size: 重叠大小
            
        Returns:
            文本块列表
        """
        chunks = []
        text_len = len(text)
        
        for i in range(0, text_len, max_size - overlap_size):
            chunk = text[i:min(i + max_size, text_len)]
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def _split_code_block(self, code_block: str) -> List[str]:
        """
        智能分割代码块，尝试保持函数和类的完整性
        
        Args:
            code_block: 代码块文本
            
        Returns:
            分割后的代码块
        """
        # 提取语言标识
        import re
        language = ""
        content = code_block
        
        match = re.match(r'```(\w*)\n(.*?)```', code_block, re.DOTALL)
        if match:
            language = match.group(1)
            content = match.group(2)
        
        # 如果代码较小，直接返回
        if len(code_block) <= settings.MAX_CHUNK_SIZE:
            return [code_block]
        
        chunks = []
        max_size = settings.MAX_CHUNK_SIZE - 8  # 为```语言和```留空间
        
        # 尝试按函数/类/模块分割
        if language in ["python", "js", "javascript", "java", "c", "cpp", "csharp", "go", "rust"]:
            # 函数/方法定义模式
            func_pattern = re.compile(r'((?:def|function|func|fn|public|private|protected|class|module|interface)\s+\w+[\s\S]*?)(?=\n\s*(?:def|function|func|fn|public|private|protected|class|module|interface)\s+|\Z)', re.DOTALL)
            
            last_end = 0
            for match in func_pattern.finditer(content):
                # 如果匹配内容太大，需要进一步分割
                if match.end() - match.start() > max_size:
                    # 将大函数/类按行分割
                    func_content = match.group(1)
                    func_chunks = self._split_by_size(func_content, max_size, 100)
                    
                    for i, chunk in enumerate(func_chunks):
                        chunks.append(f"```{language}\n{chunk}```")
                else:
                    # 处理未匹配部分
                    if match.start() > last_end:
                        unmatch = content[last_end:match.start()]
                        if unmatch.strip():
                            # 分割未匹配内容
                            unmatch_chunks = self._split_by_size(unmatch, max_size, 50)
                            for chunk in unmatch_chunks:
                                chunks.append(f"```{language}\n{chunk}```")
                    
                    # 添加函数/类定义
                    chunks.append(f"```{language}\n{match.group(1)}```")
                
                last_end = match.end()
            
            # 处理最后一部分
            if last_end < len(content):
                remaining = content[last_end:]
                if remaining.strip():
                    remaining_chunks = self._split_by_size(remaining, max_size, 50)
                    for chunk in remaining_chunks:
                        chunks.append(f"```{language}\n{chunk}```")
        else:
            # 对于不识别的语言，按行分割
            lines = content.split('\n')
            current_chunk = ""
            
            for line in lines:
                if current_chunk and len(current_chunk) + len(line) + 1 > max_size:
                    chunks.append(f"```{language}\n{current_chunk}```")
                    current_chunk = line
                else:
                    if current_chunk:
                        current_chunk += '\n' + line
                    else:
                        current_chunk = line
            
            if current_chunk:
                chunks.append(f"```{language}\n{current_chunk}```")
        
        return chunks if chunks else [code_block]
    
    def _split_list(self, list_text: str) -> List[str]:
        """
        分割列表，尝试保持列表项的完整性
        
        Args:
            list_text: 列表文本
            
        Returns:
            分割后的列表块
        """
        import re
        
        # 如果列表较小，直接返回
        if len(list_text) <= settings.MAX_CHUNK_SIZE:
            return [list_text]
        
        chunks = []
        max_size = settings.MAX_CHUNK_SIZE
        
        # 识别列表项
        list_pattern = re.compile(r'(?:^|\n)(\s*(?:[*+-]|\d+\.)\s+.*(?:\n\s+.*)*)', re.MULTILINE)
        matches = list(list_pattern.finditer(list_text))
        
        if not matches:
            # 如果无法识别列表项，按大小分割
            return self._split_by_size(list_text, max_size, 100)
        
        current_chunk = ""
        
        for match in matches:
            item = match.group(1)
            
            # 如果单个列表项就超过了最大大小
            if len(item) > max_size:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                
                # 分割大列表项
                item_chunks = self._split_by_size(item, max_size, 50)
                chunks.extend(item_chunks)
                continue
            
            # 如果当前块 + 新列表项超过大小限制
            if current_chunk and len(current_chunk) + len(item) > max_size:
                chunks.append(current_chunk)
                current_chunk = item
            else:
                # 添加到当前块
                if current_chunk:
                    current_chunk += item
                else:
                    current_chunk = item
        
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
        
    def _chunk_pdf(self, text: str) -> List[str]:
        """
        PDF文本分块
        
        Args:
            text: PDF提取的文本
            
        Returns:
            文本块列表
        """
        # PDF分块逻辑，使用通用文本分块方法
        return self._chunk_text(text)
        
    def _chunk_text(self, text: str) -> List[str]:
        """
        通用文本分块
        
        Args:
            text: 文本内容
            
        Returns:
            文本块列表
        """
        # 简单按段落分块
        chunks = []
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_size = 0
        
        for para in paragraphs:
            para_size = len(para)
            
            # 如果段落本身超过最大块大小，单独作为一个块
            if para_size > settings.MAX_CHUNK_SIZE:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                    current_size = 0
                
                # 大段落再切分
                para_chunks = self._split_large_paragraph(para)
                chunks.extend(para_chunks)
                continue
                
            # 如果加上这个段落会超过最大块大小，则开始新块
            if current_size + para_size > settings.MAX_CHUNK_SIZE:
                chunks.append(current_chunk)
                current_chunk = para
                current_size = para_size
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
                current_size += para_size
                
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks
        
    def _chunk_by_headings(self, text: str, heading_markers: List[str]) -> List[str]:
        """
        按标题分块
        
        Args:
            text: 文本内容
            heading_markers: 标题标记列表
            
        Returns:
            文本块列表
        """
        chunks = []
        lines = text.split('\n')
        
        current_chunk = ""
        current_size = 0
        
        for line in lines:
            # 检查是否为标题行
            is_heading = any(line.startswith(marker) for marker in heading_markers)
            
            # 如果是标题行且当前块非空，则开始新块
            if is_heading and current_chunk:
                chunks.append(current_chunk)
                current_chunk = line
                current_size = len(line)
            else:
                if current_chunk:
                    current_chunk += "\n" + line
                else:
                    current_chunk = line
                current_size += len(line) + 1  # +1 for newline
                
            # 如果当前块太大，也开始新块
            if current_size > settings.MAX_CHUNK_SIZE:
                chunks.append(current_chunk)
                current_chunk = ""
                current_size = 0
                
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks
        
    def _split_large_paragraph(self, paragraph: str) -> List[str]:
        """
        分割大段落
        
        Args:
            paragraph: 大段落文本
            
        Returns:
            分割后的文本块列表
        """
        chunks = []
        sentences = paragraph.replace('. ', '.\n').split('\n')
        
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # 如果句子本身超过最大块大小，按固定长度分割
            if sentence_size > settings.MAX_CHUNK_SIZE:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                    current_size = 0
                    
                # 按固定长度分割大句子
                for i in range(0, sentence_size, settings.MAX_CHUNK_SIZE):
                    chunk = sentence[i:i + settings.MAX_CHUNK_SIZE]
                    chunks.append(chunk)
                continue
                
            # 如果加上这个句子会超过最大块大小，则开始新块
            if current_size + sentence_size > settings.MAX_CHUNK_SIZE:
                chunks.append(current_chunk)
                current_chunk = sentence
                current_size = sentence_size
            else:
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                current_size += sentence_size + 1  # +1 for space
                
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks


@PipelineFactory.register_processor
class VectorizationProcessor(DocumentProcessor):
    """文档向量化处理器"""
    
    SUPPORTED_TYPES = ['pdf', 'md', 'markdown', 'docx', 'txt', 'html']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """处理文档向量化"""
        chunks = context.get('chunks', [])
        if not chunks:
            logger.warning("没有找到文本块，跳过向量化")
            return context
            
        try:
            # 这里应该调用向量化服务
            # 简化示例，实际应连接到向量数据库
            node_count = len(chunks)
            doc_id = context.get('metadata', {}).get('doc_id')
            
            # 更新上下文
            context['vectorized'] = True
            context['node_count'] = node_count
            
            logger.info(f"向量化完成，文档ID: {doc_id}, 节点数: {node_count}")
            
        except Exception as e:
            logger.error(f"向量化失败: {str(e)}")
            raise
            
        return context 


@PipelineFactory.register_processor
class HTMLProcessor(DocumentProcessor):
    """HTML文档处理器"""
    
    SUPPORTED_TYPES = ['html', 'htm']
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """处理HTML文档"""
        if context.get('file_type') not in ['html', 'htm']:
            return context
            
        # 如果已经通过MarkItDown转换过，跳过额外处理
        if context.get('markdown_content') and not context.get('markdown_conversion_failed'):
            logger.info(f"HTML文档已通过MarkItDown处理，跳过额外处理")
            return context
            
        file_path = context.get('file_path')
        try:
            # 读取HTML文件
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
                
            # 尝试提取纯文本
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                text_content = soup.get_text(separator='\n')
            except ImportError:
                # 如果没有BeautifulSoup，使用简单替换
                import re
                text_content = re.sub(r'<[^>]*>', ' ', html_content)
                text_content = re.sub(r'\s+', ' ', text_content).strip()
                
            # 更新上下文
            context['text_content'] = text_content
            
            logger.info(f"HTML文档处理完成: {file_path}")
            
        except Exception as e:
            logger.error(f"HTML处理失败: {str(e)}")
            raise
            
        return context 


@PipelineFactory.register_processor
class ParallelProcessor(DocumentProcessor):
    """并行文档处理器，用于同时处理多个文档"""
    
    def __init__(self, processor_class, max_workers=4):
        """
        初始化并行处理器
        
        Args:
            processor_class: 要并行运行的处理器类
            max_workers: 最大工作线程数
        """
        super().__init__()
        self.processor_class = processor_class
        self.max_workers = max_workers
        self.processor_instances = []
        
        # 初始化处理器实例池
        for _ in range(max_workers):
            processor = processor_class()
            self.processor_instances.append(processor)
    
    def process(self, contexts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        并行处理多个文档上下文
        
        Args:
            contexts: 文档上下文列表
            
        Returns:
            处理后的上下文列表
        """
        if not contexts:
            return []
            
        # 如果只有一个上下文，直接处理
        if len(contexts) == 1:
            processor = self.processor_instances[0]
            return [processor.process(contexts[0])]
            
        # 并行处理多个上下文
        try:
            # 使用asyncio进行并行处理
            import asyncio
            return asyncio.run(self._process_parallel(contexts))
        except Exception as e:
            logger.error(f"并行处理失败: {str(e)}")
            # 回退到串行处理
            results = []
            for context in contexts:
                processor = self.processor_instances[0]
                results.append(processor.process(context))
            return results
    
    async def _process_parallel(self, contexts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        使用异步方式并行处理多个上下文
        
        Args:
            contexts: 文档上下文列表
            
        Returns:
            处理后的上下文列表
        """
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        
        # 创建线程池
        executor = ThreadPoolExecutor(max_workers=self.max_workers)
        loop = asyncio.get_event_loop()
        
        # 创建任务
        tasks = []
        for i, context in enumerate(contexts):
            # 选择处理器实例
            processor_index = i % len(self.processor_instances)
            processor = self.processor_instances[processor_index]
            
            # 创建异步任务
            task = loop.run_in_executor(
                executor,
                processor.process,
                context
            )
            tasks.append(task)
        
        # 等待所有任务完成
        results = await asyncio.gather(*tasks)
        
        return results


class AsyncDocumentPipeline:
    """异步文档处理管道，支持并行处理多个文档"""
    
    def __init__(self, processors=None, max_workers=4):
        """
        初始化异步文档处理管道
        
        Args:
            processors: 处理器列表
            max_workers: 最大工作线程数
        """
        self.processors = processors or []
        self.max_workers = max_workers
    
    async def process_documents(self, files: List[str]) -> List[Dict[str, Any]]:
        """
        异步处理多个文档
        
        Args:
            files: 文件路径列表
            
        Returns:
            处理结果列表
        """
        if not files:
            return []
            
        # 为每个文件创建初始上下文
        contexts = []
        for file_path in files:
            # 提取文件类型
            file_type = self._get_file_type(file_path)
            context = {
                'file_path': file_path,
                'file_type': file_type,
                'metadata': {
                    'doc_id': self._generate_doc_id(file_path),
                    'filename': os.path.basename(file_path),
                    'created_at': datetime.datetime.now().isoformat()
                }
            }
            contexts.append(context)
        
        # 逐步应用每个处理器
        for processor in self.processors:
            # 检查是否支持批处理
            if isinstance(processor, ParallelProcessor):
                # 并行处理
                contexts = processor.process(contexts)
            else:
                # 创建处理器的并行包装
                parallel_processor = ParallelProcessor(
                    processor.__class__, 
                    max_workers=self.max_workers
                )
                contexts = parallel_processor.process(contexts)
        
        return contexts
    
    def _get_file_type(self, file_path: str) -> str:
        """从文件路径获取文件类型"""
        _, ext = os.path.splitext(file_path)
        return ext.lstrip('.').lower()
    
    def _generate_doc_id(self, file_path: str) -> str:
        """为文件生成唯一ID"""
        import hashlib
        import time
        
        # 使用文件路径和时间戳生成唯一ID
        unique_str = f"{file_path}_{time.time()}"
        doc_id = hashlib.md5(unique_str.encode()).hexdigest()
        
        return doc_id


# 扩展ChunkingProcessor以支持增量更新
@PipelineFactory.register_processor
class IncrementalChunkingProcessor(ChunkingProcessor):
    """支持增量更新的文档分块处理器"""
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        处理文档分块，支持增量更新
        
        如果提供了旧内容，则只处理变更的部分
        """
        # 检查是否是增量更新模式
        is_incremental = context.get('incremental_update', False)
        old_content = context.get('old_content', '')
        
        if is_incremental and old_content:
            # 执行增量更新
            new_content = context.get('text_content', '')
            if not new_content:
                return context
            
            # 对比新旧内容，仅处理变更部分
            chunks = self._incremental_document_update(
                new_content, 
                old_content,
                context.get('file_type', '')
            )
            
            # 更新上下文
            context['chunks'] = chunks
            context['chunk_count'] = len(chunks)
            context['incremental_processed'] = True
            
            logger.info(f"增量文档分块完成，共 {len(chunks)} 个块")
            
        else:
            # 执行常规分块
            context = super().process(context)
            
        return context
    
    def _incremental_document_update(self, new_content: str, old_content: str, file_type: str) -> List[str]:
        """
        增量更新文档，只处理变更部分
        
        Args:
            new_content: 新文档内容
            old_content: 旧文档内容
            file_type: 文件类型
            
        Returns:
            更新后的分块列表
        """
        import difflib
        
        # 计算差异
        diff = list(difflib.ndiff(old_content.splitlines(), new_content.splitlines()))
        
        # 提取变更的段落
        changes = []
        current_section = []
        in_change = False
        
        for line in diff:
            if line.startswith('- '):
                # 删除行，标记为变更
                in_change = True
            elif line.startswith('+ '):
                # 添加行，标记为变更并保存内容
                in_change = True
                current_section.append(line[2:])
            elif line.startswith('  '):
                # 未变更行
                if in_change:
                    # 如果之前有变更，添加上下文
                    current_section.append(line[2:])
                    
                    # 如果上下文足够长，保存当前变更段落
                    if len(current_section) >= 5:
                        changes.append('\n'.join(current_section))
                        current_section = []
                        in_change = False
        
        # 处理最后一个变更段落
        if current_section:
            changes.append('\n'.join(current_section))
        
        # 如果没有检测到变更，返回空列表
        if not changes:
            logger.info("未检测到文档变更，跳过处理")
            return []
        
        # 处理变更的段落
        chunks = []
        for section in changes:
            # 根据文件类型选择合适的分块策略
            if file_type in ['md', 'markdown']:
                section_chunks = self._chunk_markdown(section)
            elif file_type == 'pdf':
                section_chunks = self._chunk_pdf(section)
            else:
                section_chunks = self._chunk_text(section)
            
            chunks.extend(section_chunks)
        
        return chunks 


@PipelineFactory.register_processor
class ContextCompressorProcessor(DocumentProcessor):
    """
    上下文压缩处理器 - 基于LangChain的设计思想
    
    根据查询内容动态压缩文档块，提高检索效率
    """
    
    def __init__(self, embedding_model=None, compression_ratio=0.7):
        """
        初始化上下文压缩处理器
        
        Args:
            embedding_model: 用于计算相关性的嵌入模型，如果为None则使用简单的关键词匹配
            compression_ratio: 压缩比例，默认压缩到原始大小的70%
        """
        super().__init__()
        self.embedding_model = embedding_model
        self.compression_ratio = compression_ratio
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        处理文档上下文，如果提供了查询，则进行压缩
        
        Args:
            context: 文档上下文
            
        Returns:
            处理后的上下文
        """
        chunks = context.get('chunks', [])
        query = context.get('query', '')
        
        # 如果没有块或查询，直接返回
        if not chunks or not query:
            return context
        
        # 压缩文档块
        compressed_chunks = self._compress_context(chunks, query)
        
        # 更新上下文
        context['original_chunks'] = chunks
        context['compressed_chunks'] = compressed_chunks
        context['chunks'] = compressed_chunks  # 替换为压缩后的块
        context['compression_applied'] = True
        context['compression_ratio'] = self.compression_ratio
        
        logger.info(f"上下文压缩完成: 原始 {len(chunks)} 块，压缩后 {len(compressed_chunks)} 块")
        
        return context
    
    def _compress_context(self, chunks: List[str], query: str) -> List[str]:
        """
        智能上下文压缩 - 类似LangChain的ContextualizationDocumentCompressor
        
        根据查询动态调整文本块的大小，保留与查询相关的部分
        
        Args:
            chunks: 原始文档块列表
            query: 查询内容
            
        Returns:
            压缩后的文档块列表
        """
        compressed_chunks = []
        
        for chunk in chunks:
            # 计算块与查询的相关性分数
            relevance_score = self._calculate_relevance(chunk, query)
            
            # 根据相关性分数决定压缩程度
            if relevance_score > 0.8:
                # 高相关，保留完整块
                compressed_chunks.append(chunk)
            elif relevance_score > 0.5:
                # 中等相关，提取核心段落
                core = self._extract_core_sentences(chunk, query)
                compressed_chunks.append(core)
            else:
                # 低相关，只保留摘要
                summary = self._generate_summary(chunk)
                if len(summary) > 50:  # 确保摘要有意义
                    compressed_chunks.append(summary)
        
        return compressed_chunks
    
    def _calculate_relevance(self, text: str, query: str) -> float:
        """
        计算文本与查询的相关性分数
        
        Args:
            text: 文本内容
            query: 查询内容
            
        Returns:
            相关性分数，范围[0,1]
        """
        # 如果有嵌入模型，使用语义相似度
        if self.embedding_model:
            try:
                text_embedding = self.embedding_model.encode(text)
                query_embedding = self.embedding_model.encode(query)
                
                # 计算余弦相似度
                similarity = self._cosine_similarity(text_embedding, query_embedding)
                return float(similarity)
            except Exception as e:
                logger.warning(f"使用嵌入模型计算相关性失败: {str(e)}，回退到关键词匹配")
        
        # 回退到简单的关键词匹配
        return self._keyword_relevance(text, query)
    
    def _cosine_similarity(self, vec1, vec2):
        """计算余弦相似度"""
        import numpy as np
        dot_product = np.dot(vec1, vec2)
        norm_a = np.linalg.norm(vec1)
        norm_b = np.linalg.norm(vec2)
        
        if norm_a == 0 or norm_b == 0:
            return 0
            
        return dot_product / (norm_a * norm_b)
    
    def _keyword_relevance(self, text: str, query: str) -> float:
        """
        基于关键词匹配计算相关性
        
        Args:
            text: 文本内容
            query: 查询内容
            
        Returns:
            相关性分数
        """
        # 简化版的关键词匹配
        text_lower = text.lower()
        query_terms = [t for t in query.lower().split() if len(t) > 2]
        
        if not query_terms:
            return 0.5  # 默认中等相关
        
        # 计算匹配的词数
        matches = sum(1 for term in query_terms if term in text_lower)
        relevance = matches / len(query_terms)
        
        return min(1.0, relevance)  # 确保不超过1.0
    
    def _extract_core_sentences(self, text: str, query: str) -> str:
        """
        提取与查询相关的核心句子
        
        Args:
            text: 文本内容
            query: 查询内容
            
        Returns:
            核心句子组成的文本
        """
        import re
        
        # 分割成句子
        sentences = re.split(r'(?<=[.!?。！？])\s+', text)
        
        # 如果句子太少，直接返回原文
        if len(sentences) <= 3:
            return text
        
        # 计算每个句子的相关性得分
        sentence_scores = []
        for sentence in sentences:
            score = self._keyword_relevance(sentence, query)
            sentence_scores.append((sentence, score))
        
        # 排序并选择得分最高的前N个句子
        sentence_scores.sort(key=lambda x: x[1], reverse=True)
        
        # 确定要保留的句子数，压缩到指定比例
        keep_count = max(1, int(len(sentences) * self.compression_ratio))
        top_sentences = [s[0] for s in sentence_scores[:keep_count]]
        
        # 恢复原始顺序
        ordered_sentences = [s for s in sentences if s in top_sentences]
        
        return " ".join(ordered_sentences)
    
    def _generate_summary(self, text: str) -> str:
        """
        为文本生成简短摘要
        
        Args:
            text: 文本内容
            
        Returns:
            文本摘要
        """
        # 简单摘要：取前几个句子
        import re
        sentences = re.split(r'(?<=[.!?。！？])\s+', text)
        
        # 取前2-3个句子作为摘要
        summary_length = min(3, len(sentences))
        summary = " ".join(sentences[:summary_length])
        
        # 如果摘要太长，截断
        if len(summary) > 200:
            summary = summary[:197] + "..."
            
        return summary 


@PipelineFactory.register_processor
class MetadataAwareChunkingProcessor(ChunkingProcessor):
    """
    元数据感知的文档分块处理器 - 基于RAGFlow的设计思想
    
    从文档内容中提取元数据，并与分块关联，增强检索精度
    """
    
    def process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        处理文档分块并提取元数据
        
        Args:
            context: 文档上下文
            
        Returns:
            处理后的上下文
        """
        # 使用父类的分块方法
        context = super().process(context)
        
        chunks = context.get('chunks', [])
        if not chunks:
            return context
            
        # 提取文档整体元数据
        doc_metadata = self._extract_document_metadata(context)
        
        # 为每个块提取和关联元数据
        chunk_metadata_list = []
        enhanced_chunks = []
        
        for i, chunk in enumerate(chunks):
            # 提取块级元数据
            chunk_metadata = self._extract_metadata_from_content(chunk)
            
            # 合并文档元数据和块级元数据
            combined_metadata = {**doc_metadata, **chunk_metadata}
            
            # 为检索系统保存元数据
            chunk_metadata_list.append(combined_metadata)
            
            # 可选：在块内容中添加元数据标记（用于调试或特殊处理）
            if context.get('embed_metadata_in_chunks', False):
                enhanced_chunk = self._embed_metadata_in_chunk(chunk, combined_metadata)
                enhanced_chunks.append(enhanced_chunk)
        
        # 更新上下文
        context['chunk_metadata'] = chunk_metadata_list
        if enhanced_chunks:
            context['original_chunks'] = chunks
            context['chunks'] = enhanced_chunks
        
        context['metadata_enhanced'] = True
        
        logger.info(f"文档分块元数据增强完成: {len(chunks)} 个块")
        
        return context
    
    def _extract_document_metadata(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        从文档上下文中提取文档级元数据
        
        Args:
            context: 文档上下文
            
        Returns:
            文档元数据字典
        """
        metadata = {}
        
        # 基本文件信息
        metadata['doc_id'] = context.get('metadata', {}).get('doc_id', '')
        metadata['filename'] = context.get('metadata', {}).get('filename', '')
        metadata['file_type'] = context.get('file_type', '')
        metadata['file_size'] = context.get('file_size', 0)
        
        # 获取文档创建和修改时间
        file_path = context.get('file_path', '')
        if file_path and os.path.exists(file_path):
            import datetime
            metadata['created_time'] = datetime.datetime.fromtimestamp(
                os.path.getctime(file_path)
            ).isoformat()
            metadata['modified_time'] = datetime.datetime.fromtimestamp(
                os.path.getmtime(file_path)
            ).isoformat()
        
        # 文档结构信息
        if 'document_structure' in context:
            structure = context['document_structure']
            
            # 提取标题信息
            if 'headings' in structure:
                # 提取主标题
                main_headings = [h[1] for h in structure['headings'] if h[0] == 1]
                if main_headings:
                    metadata['title'] = main_headings[0]
                
                # 所有标题列表
                metadata['headings'] = [h[1] for h in structure['headings']]
            
            # 提取图片信息
            if 'images' in structure:
                metadata['has_images'] = len(structure['images']) > 0
                metadata['image_count'] = len(structure['images'])
            
            # 提取代码块信息
            if 'code_blocks' in structure:
                metadata['has_code'] = len(structure['code_blocks']) > 0
                metadata['code_block_count'] = len(structure['code_blocks'])
                
                # 提取编程语言
                languages = set()
                for lang, _ in structure['code_blocks']:
                    if lang:
                        languages.add(lang)
                if languages:
                    metadata['programming_languages'] = list(languages)
        
        # 处理文档引用关系
        if context.get('markdown_content'):
            references = self._extract_references(context['markdown_content'])
            if references:
                metadata['references'] = references
        
        return metadata
    
    def _extract_metadata_from_content(self, content: str) -> Dict[str, Any]:
        """
        从文本内容中提取元数据
        
        Args:
            content: 文本内容
            
        Returns:
            元数据字典
        """
        metadata = {
            'entities': [],
            'keywords': [],
            'topic_areas': [],
            'time_references': []
        }
        
        # 提取关键词
        try:
            import jieba.analyse
            keywords = jieba.analyse.extract_tags(content, topK=10)
            metadata['keywords'] = keywords
        except ImportError:
            # 如果没有jieba，使用简单的词频统计
            metadata['keywords'] = self._extract_keywords_by_frequency(content)
        
        # 提取时间引用
        import re
        time_pattern = re.compile(r'\d{4}年|\d{4}-\d{2}|\d{4}/\d{2}|\d{4}\.\d{2}')
        time_refs = time_pattern.findall(content)
        if time_refs:
            metadata['time_references'] = time_refs
        
        # 提取实体（人名、地点、组织等）
        entities = self._extract_entities(content)
        if entities:
            metadata['entities'] = entities
        
        # 尝试确定主题领域
        topics = self._identify_topic_areas(content, metadata['keywords'])
        if topics:
            metadata['topic_areas'] = topics
        
        return metadata
    
    def _extract_keywords_by_frequency(self, text: str, top_k=10) -> List[str]:
        """
        通过词频提取关键词
        
        Args:
            text: 文本内容
            top_k: 返回的关键词数量
            
        Returns:
            关键词列表
        """
        import re
        from collections import Counter
        
        # 简单分词（中英文混合）
        words = re.findall(r'[\w\u4e00-\u9fff]+', text.lower())
        
        # 过滤停用词（简化版）
        stopwords = {'的', '了', '和', '是', '在', '有', '与', '这', '那', '之', '或', '及', 'a', 'an', 'the', 'and', 'or', 'in', 'on', 'at', 'to', 'for', 'with'}
        words = [w for w in words if w not in stopwords and len(w) > 1]
        
        # 计算词频
        word_freq = Counter(words)
        
        # 返回top_k关键词
        return [w for w, _ in word_freq.most_common(top_k)]
    
    def _extract_entities(self, text: str) -> List[str]:
        """
        提取实体
        
        Args:
            text: 文本内容
            
        Returns:
            实体列表
        """
        entities = []
        
        # 实现简单的规则匹配
        # 例如：识别大写开头的词作为潜在实体（英文）
        import re
        
        # 英文实体识别
        eng_entity_pattern = re.compile(r'\b[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*)*\b')
        eng_entities = eng_entity_pattern.findall(text)
        
        # 过滤单个词且太常见的词
        common_words = {'I', 'You', 'He', 'She', 'It', 'We', 'They', 'The', 'A', 'An'}
        eng_entities = [e for e in eng_entities if e not in common_words and (len(e.split()) > 1 or len(e) > 3)]
        
        entities.extend(eng_entities)
        
        # 中文人名识别（简易版）
        cn_name_pattern = re.compile(r'[\u4e00-\u9fff]{2,3}(?:先生|女士|老师|教授|博士)')
        cn_names = cn_name_pattern.findall(text)
        
        if cn_names:
            # 去掉称谓后缀
            cn_names = [re.sub(r'(先生|女士|老师|教授|博士)$', '', name) for name in cn_names]
            entities.extend(cn_names)
        
        return list(set(entities))  # 去重
    
    def _identify_topic_areas(self, text: str, keywords: List[str]) -> List[str]:
        """
        识别主题领域
        
        Args:
            text: 文本内容
            keywords: 已提取的关键词
            
        Returns:
            主题领域列表
        """
        # 根据关键词判断主题领域
        topics = set()
        
        # 技术主题
        tech_terms = {'python', 'java', 'javascript', 'html', 'css', 'coding', 'algorithm', 'programming', 
                       '编程', '算法', '代码', '开发', '软件', '数据库', '测试', '部署', '服务器'}
        if any(kw.lower() in tech_terms for kw in keywords):
            topics.add('技术')
        
        # 商业主题
        business_terms = {'market', 'business', 'company', 'strategy', 'investment', 'customer', 'product',
                          '市场', '商业', '公司', '策略', '投资', '客户', '产品', '营销', '销售'}
        if any(kw.lower() in business_terms for kw in keywords):
            topics.add('商业')
        
        # 教育主题
        education_terms = {'education', 'teaching', 'learning', 'student', 'school', 'university', 'course',
                           '教育', '教学', '学习', '学生', '学校', '大学', '课程', '培训'}
        if any(kw.lower() in education_terms for kw in keywords):
            topics.add('教育')
        
        # 医疗主题
        medical_terms = {'medical', 'healthcare', 'disease', 'patient', 'treatment', 'medicine', 'doctor',
                         '医疗', '健康', '疾病', '患者', '治疗', '药物', '医生', '医院'}
        if any(kw.lower() in medical_terms for kw in keywords):
            topics.add('医疗')
        
        # 扩展更多领域...
        
        return list(topics)
    
    def _extract_references(self, text: str) -> List[Dict[str, str]]:
        """
        提取文本中的引用
        
        Args:
            text: 文本内容
            
        Returns:
            引用列表，每个引用为一个字典
        """
        references = []
        
        # 提取URL
        import re
        url_pattern = re.compile(r'https?://\S+')
        urls = url_pattern.findall(text)
        
        for url in urls:
            references.append({
                'type': 'url',
                'value': url
            })
        
        # 提取文献引用 [1], [2] 等
        citation_pattern = re.compile(r'\[(\d+)\]')
        citations = citation_pattern.findall(text)
        
        for citation in citations:
            references.append({
                'type': 'citation',
                'value': citation
            })
        
        return references
    
    def _embed_metadata_in_chunk(self, chunk: str, metadata: Dict[str, Any]) -> str:
        """
        在块内容中嵌入元数据（用于调试或特殊处理）
        
        Args:
            chunk: 块内容
            metadata: 元数据
            
        Returns:
            嵌入元数据的块内容
        """
        # 选择重要的元数据字段
        important_fields = ['keywords', 'entities', 'topic_areas']
        
        metadata_text = "\n\n---\n元数据: "
        for field in important_fields:
            if field in metadata and metadata[field]:
                if isinstance(metadata[field], list):
                    metadata_text += f"{field}=[{', '.join(metadata[field])}]; "
                else:
                    metadata_text += f"{field}={metadata[field]}; "
        
        return chunk + metadata_text